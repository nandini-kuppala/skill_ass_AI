# !pip install crewai python-dotenv PyPDF2 python-docx PyMuPDF langchain langchain-google-genai google-generativeai beautifulsoup4 requests
import os
import json
import re
import traceback
import PyPDF2
import docx
import fitz

# For web scraping
import requests
# For LLM integrations
import google.generativeai as genai
from langchain_community.chat_models import ChatLiteLLM
from crewai import LLM
# For Crew AI
from crewai import Agent, Task, Process, Crew
from crewai.tools.base_tool import BaseTool as Tool
# Load environment variables
import streamlit as st

SERPER_API_KEY = st.secrets["SERPER_API_KEY"]
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

# Verify API keys are loaded
print(f"SERPER_API_KEY loaded: {'Yes' if SERPER_API_KEY else 'No'}")
print(f"GEMINI_API_KEY loaded: {'Yes' if GEMINI_API_KEY else 'No'}")
# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)
gemini_pro = genai.GenerativeModel('gemini-2.0-flash-lite')
print("Environment and API configurations loaded successfully.")



class SearchTool(Tool):
    name: str = "Search Tool"
    description: str = "Useful for searching information about technologies, companies, or people online."
    
    def _run(self, query: str) -> str:  # Added self parameter
        try:
            url = "https://google.serper.dev/search"
            payload = json.dumps({"q": query})
            headers = {
                'X-API-KEY': SERPER_API_KEY,
                'Content-Type': 'application/json'
            }
            response = requests.request("POST", url, headers=headers, data=payload)
            return response.text
        except Exception as e:
            return f"Error during search: {str(e)}"

# 4.2 GitHub Profile Scraper
class GitHubProfileScraper(Tool):
    name: str = "GitHub Profile Scraper"
    description: str = "Scrapes a GitHub profile to gather information about repositories, contributions, and skills."
    
    def _run(self, github_username: str) -> str:  # Added self parameter
        try:
            # Rest of the method stays the same
            profile_url = f"https://api.github.com/users/{github_username}"
            repos_url = f"https://api.github.com/users/{github_username}/repos"
            
            # Get profile information
            profile_response = requests.get(profile_url)

            if profile_response.status_code != 200:
                return f"Error: Could not retrieve GitHub profile for {github_username}."
            
            profile_data = profile_response.json()
            
            # Get repositories information
            repos_response = requests.get(repos_url)
            if repos_response.status_code != 200:
                return f"Error: Could not retrieve repositories for {github_username}."
            
            repos_data = repos_response.json()
            
            # Extract relevant information
            total_repos = len(repos_data)
            languages_used = {}
            starred_repos = 0
            forks = 0
            
            for repo in repos_data:
                forks += repo['forks_count']
                starred_repos += repo['stargazers_count']
                
                # Get languages used in each repository
                if not repo['fork']:  # Only consider non-forked repos
                    lang_url = repo['languages_url']
                    lang_response = requests.get(lang_url)
                    if lang_response.status_code == 200:
                        langs = lang_response.json()
                        for lang, bytes_of_code in langs.items():
                            if lang in languages_used:
                                languages_used[lang] += bytes_of_code
                            else:
                                languages_used[lang] = bytes_of_code
            
            # Sort languages by bytes of code
            sorted_languages = sorted(languages_used.items(), key=lambda x: x[1], reverse=True)
            top_languages = [lang for lang, _ in sorted_languages[:5]]
            
            # Format the results
            result = {
                "username": github_username,
                "name": profile_data.get('name', 'Not available'),
                "bio": profile_data.get('bio', 'Not available'),
                "public_repos": total_repos,
                "followers": profile_data.get('followers', 0),
                "following": profile_data.get('following', 0),
                "top_languages": top_languages,
                "starred_repos_count": starred_repos,
                "forks_count": forks,
                "profile_created_at": profile_data.get('created_at', 'Not available')
            }
            
            return json.dumps(result, indent=2)
        except Exception as e:
            return f"Error scraping GitHub profile: {str(e)}\n{traceback.format_exc()}"

# 4.3 LeetCode Profile Scraper
class LeetCodeProfileScraper(Tool):
    name: str = "LeetCode Profile Scraper"
    description: str = "Scrapes a LeetCode profile to gather information about solved problems and contest ratings."
    
    def _run(self, leetcode_username: str) -> str:  # Added self parameter
        try:
            # Using public GraphQL API for LeetCode
            url = "https://leetcode.com/graphql"
            
            # Query to get user profile information
            query = """
            query userPublicProfile($username: String!) {
              matchedUser(username: $username) {
                username
                submitStats: submitStatsGlobal {
                  acSubmissionNum {
                    difficulty
                    count
                    submissions
                  }
                }
                profile {
                  ranking
                  reputation
                  starRating
                }
                badges {
                  id
                  name
                  icon
                }
              }
            }
            """
            
            variables = {"username": leetcode_username}
            payload = {"query": query, "variables": variables}
            headers = {
                "Content-Type": "application/json",
                "Referer": f"https://leetcode.com/{leetcode_username}/"
            }
            
            response = requests.post(url, headers=headers, json=payload)
            
            if response.status_code != 200:
                return f"Error: Could not retrieve LeetCode profile for {leetcode_username}."
            
            data = response.json()
            
            if not data.get('data', {}).get('matchedUser'):
                return f"Error: LeetCode user {leetcode_username} not found."
            
            user_data = data['data']['matchedUser']
            
            # Extract relevant information
            submission_stats = user_data['submitStats']['acSubmissionNum']
            problem_counts = {}
            
            for stat in submission_stats:
                problem_counts[stat['difficulty']] = stat['count']
            
            total_solved = sum(problem_counts.values())
            
            # Format the results
            result = {
                "username": leetcode_username,
                "total_problems_solved": total_solved,
                "problems_by_difficulty": problem_counts,
                "ranking": user_data['profile'].get('ranking', 'Not available'),
                "reputation": user_data['profile'].get('reputation', 'Not available'),
                "badges": [badge['name'] for badge in user_data.get('badges', [])]
            }
            
            return json.dumps(result, indent=2)
        except Exception as e:
            return f"Error scraping LeetCode profile: {str(e)}\n{traceback.format_exc()}"

def extract_text_from_pdf(pdf_file):
        """Extract text from a PDF file using multiple extraction methods for redundancy."""
        text = ""
        temp_file_path = None
        
        try:
            # First attempt: Use PyMuPDF (fitz) for better text extraction
            pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
            
            # Extract text page by page
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_text = page.get_text()
                text += page_text
                
                # Extract links from annotations
                links = page.get_links()
                for link in links:
                    if 'uri' in link:
                        text += f"\nLink: {link['uri']}\n"
            
            pdf_document.close()
            
            # If text extraction yielded minimal results, try OCR fallback
            if len(text.strip()) < 100:
                print("Minimal text extracted. PDF might be scanned or secured.")
                # Here you could implement OCR using pytesseract
                pass
                
        except Exception as e:
            print(f"Error extracting text with PyMuPDF: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                pdf_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text
            except Exception as e2:
                print(f"Error with fallback PDF extraction: {str(e2)}")
                
        return text
def extract_text_from_docx(docx_file):
        
        """Extract text from a DOCX file including hyperlinks."""
        text = ""
        try:
            doc = docx.Document(docx_file)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
                
                # Extract hyperlinks from paragraph runs
                for run in para.runs:
                    if run.element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        for hyperlink in run.element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            for relationship_id in hyperlink.attrib.values():
                                # This is a simple approach - a more robust solution would look up the relationship ID
                                # in the document's relationships to get the actual URL
                                text += f"\nHyperlink: {relationship_id}\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
                    
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            
        return text

def extract_text_from_document(file):
        """Extract text from various document formats with enhanced link extraction."""
        try:
            # Get file extension
            file_name = file.name.lower() if hasattr(file, 'name') else "unknown_file"
            
            if file_name.endswith('.pdf'):
                return extract_text_from_pdf(file)
            elif file_name.endswith('.docx'):
                return extract_text_from_docx(file)
            elif file_name.endswith('.txt'):
                return file.read().decode('utf-8')
            else:
                return f"Unsupported file format: {file_name.split('.')[-1]}"
                
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return f"Error processing document: {str(e)}"


def extract_profile_links(text):
    """Extract coding profile links and other professional profile links from text."""
    profiles = {}
    
    # GitHub profile patterns - enhanced with multiple formats
    github_patterns = [
        r'github\.com/([a-zA-Z0-9_-]+)',
        r'github:?\s*:?\s*([a-zA-Z0-9_-]+)',
        r'GitHub\s*Profile:?\s*([a-zA-Z0-9_-]+)',
        r'[Gg][Ii][Tt][Hh][Uu][Bb][:\s]+\s*([a-zA-Z0-9_-]+)',
        r'[Gg][Ii][Tt][Hh][Uu][Bb].*?[:/]([a-zA-Z0-9_-]+)',
    ]
    
    # LeetCode profile patterns - enhanced with multiple formats
    leetcode_patterns = [
        r'leetcode\.com/([a-zA-Z0-9_-]+)',
        r'leetcode:?\s*:?\s*([a-zA-Z0-9_-]+)',
        r'LeetCode\s*Profile:?\s*([a-zA-Z0-9_-]+)',
        r'[Ll][Ee][Ee][Tt][Cc][Oo][Dd][Ee][:\s]+\s*([a-zA-Z0-9_-]+)',
        r'[Ll][Ee][Ee][Tt][Cc][Oo][Dd][Ee].*?[:/]([a-zA-Z0-9_-]+)',
    ]
    
    # Certification link patterns
    certification_patterns = [
        r'certification:?\s*[^\n]*?(https?://[^\s]+)',
        r'certificate:?\s*[^\n]*?(https?://[^\s]+)',
        r'(https?://[^\s]*certificate[^\s]*)',
        r'(https?://[^\s]*certify[^\s]*)',
        r'(https?://[^\s]*verify[^\s]*)',
        r'(https?://[^\s]*credential[^\s]*)',
    ]
    
    # Project link patterns
    project_patterns = [
        r'project:?\s*[^\n]*?(https?://[^\s]+)',
        r'(https?://github\.com/[^\s]+/[^\s]+)',
        r'(https?://gitlab\.com/[^\s]+/[^\s]+)',
        r'(https?://bitbucket\.org/[^\s]+/[^\s]+)',
    ]
    
    # Extract GitHub username
    for pattern in github_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Clean up the username (remove trailing punctuation, etc.)
            username = re.sub(r'[^a-zA-Z0-9_-]', '', matches[0])
            profiles['github'] = username
            break
    
    # Extract LeetCode username
    for pattern in leetcode_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Clean up the username
            username = re.sub(r'[^a-zA-Z0-9_-]', '', matches[0])
            profiles['leetcode'] = username
            break
    
    # Extract certification links
    certification_links = []
    for pattern in certification_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            # Clean up the URL
            url = match.rstrip(',.;:()[]{}"\'')
            certification_links.append(url)
    
    if certification_links:
        profiles['certifications'] = certification_links
    
    # Extract project links
    project_links = []
    for pattern in project_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            # Clean up the URL
            url = match.rstrip(',.;:()[]{}"\'')
            project_links.append(url)
    
    if project_links:
        profiles['projects'] = project_links
    
    return profiles

def extract_certifications(text):
    """Extract certification information from text."""
    certifications = []
    
    # Look for certification sections
    cert_section_patterns = [
        r'(?:CERTIFICATIONS?|CERTIFICATES?|QUALIFICATIONS?)[:\s]*(.*?)(?:EDUCATION|EXPERIENCE|SKILLS|PROJECTS|\Z)',
        r'(?:CERTIFICATIONS?|CERTIFICATES?)[^\n]*\n(.*?)(?:\n\s*\n|\Z)',
    ]
    
    cert_text = ""
    for pattern in cert_section_patterns:
        matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
        if matches:
            cert_text = matches[0]
            break
    
    if not cert_text:
        # Try to find individual certifications if no section was found
        individual_cert_pattern = r'(?:certified|certification|certificate)[\s:]+([^\n,]+)'
        matches = re.findall(individual_cert_pattern, text, re.IGNORECASE)
        for match in matches:
            certifications.append({
                "name": match.strip(),
                "issuer": "Unknown",
                "date": "Unknown",
                "link": None
            })
        return certifications
    
    # Split certification text into individual certifications
    cert_entries = re.split(r'\n+', cert_text)
    for entry in cert_entries:
        if len(entry.strip()) < 5:  # Skip very short lines
            continue
            
        cert = {
            "name": entry.strip(),
            "issuer": "Unknown",
            "date": "Unknown",
            "link": None
        }
        
        # Try to extract issuer
        issuer_match = re.search(r'(?:from|by|issued by|through)\s+([^,\n]+)', entry, re.IGNORECASE)
        if issuer_match:
            cert["issuer"] = issuer_match.group(1).strip()
        
        # Try to extract date
        date_match = re.search(r'(?:issued|completed|received)?\s*(?:in|on)?\s*(\d{1,2}/\d{1,2}/\d{2,4}|\w+\s+\d{4}|\d{4})', entry, re.IGNORECASE)
        if date_match:
            cert["date"] = date_match.group(1).strip()
        
        # Try to extract link
        link_match = re.search(r'(https?://[^\s]+)', entry)
        if link_match:
            cert["link"] = link_match.group(1).strip(',.;:()"\'[]{}')
        
        certifications.append(cert)
    
    return certifications
def create_document_parser_agent():
    """Create an improved document parser agent with enhanced profile extraction capabilities."""
    llm = LLM(
    model="gemini/gemini-2.0-flash-lite",  
    temperature=0.2,
    api_key=st.secrets["GEMINI_API_KEY"] 
    )
    return Agent(
        role="Document Parser Specialist",
        goal="Extract comprehensive structured information from candidate resumes with a focus on profile links",
        backstory="""You are an AI specialist in parsing and extracting structured information from resume documents.
        Your expertise allows you to identify key elements like personal information, skills, experience, education,
        and public coding profiles from various document formats. You have a particular talent for finding and
        extracting URLs and profile links that may be embedded in the document, even when they are not explicitly
        labeled. You understand various resume formats and can identify GitHub, LeetCode, and certification links
        regardless of how they're presented.""",
        verbose=True,
        allow_delegation=True,
        tools=[],
        llm=llm
    )

def create_profile_scraper_agent():
    """Create an improved profile scraper agent."""
    llm = LLM(
    model="gemini/gemini-2.0-flash-lite",  
    temperature=0.2,
    api_key=st.secrets["GEMINI_API_KEY"] 
    )
    return Agent(
        role="Profile Data Collector",
        goal="Gather comprehensive data from candidates' public coding profiles",
        backstory="""You specialize in collecting data from coding platforms like GitHub and LeetCode.
        Your work provides evidence of candidates' technical abilities beyond what they claim in resumes.
        You thoroughly analyze repositories, contributions, and problem-solving patterns to assess real-world
        coding abilities. You know how to extract information about languages used, project complexity,
        and development patterns from code repositories. For competitive programming profiles, you identify
        problem-solving skills and algorithmic thinking capabilities.""",
        verbose=True,
        allow_delegation=True,
        tools=[
            GitHubProfileScraper(),
            LeetCodeProfileScraper(),
            SearchTool()
        ],
        llm=llm
    )


def create_technical_evaluator_agent():
    """Create an improved technical evaluator agent."""
    llm = LLM(
    model="gemini/gemini-2.0-flash-lite",  
    temperature=0.2,
    api_key=st.secrets["GEMINI_API_KEY"] 
    )    
    return Agent(
        role="Technical Skills Evaluator",
        goal="Assess technical proficiency based on resume claims, coding profiles, and certification evidence",
        backstory="""You are an experienced technical evaluator with deep knowledge across programming languages, frameworks,
        and software development practices. You can assess a candidate's real technical abilities by analyzing their coding profiles such as leetcode profile, 
        coding projects and contributions, contributions on github, internship experience and contribution in the projects, and cource certifications, acadamic achievements, Publications, Participation in hackathons . You understand the
        difference between claimed skills and demonstrated expertise, and can identify when a candidate has practical
        experience versus theoretical knowledge. You evaluate technical depth, project complexity and novelty and the candidates role in the projects""",
        verbose=True,
        allow_delegation=True,
        tools=[],
        llm=llm
    )

def create_job_alignment_agent():
    """Create an improved job alignment agent."""
    llm = LLM(
    model="gemini/gemini-2.0-flash-lite",  
    temperature=0.2,
    api_key=st.secrets["GEMINI_API_KEY"] 
    )
    return Agent(
        role="Job Requirements Specialist",
        goal="Determine how well a candidate matches the specific technical requirements of a job",
        backstory="""With your expertise in technical recruitment, you excel at mapping candidate skills against job
        requirements. You understand both what companies need and how candidate abilities translate to on-the-job performance.
        You can identify both direct skill matches and transferable skills that might not be explicitly mentioned but are
        relevant to the role. You understand the difference between critical requirements and nice-to-have qualifications,
        weighing them appropriately in your assessment. You know how to evaluate cultural fit indicators and growth potential
        based on career progression and learning patterns shown in the candidate's profile.""",
        verbose=True,
        allow_delegation=True,
        tools=[],
        llm=llm
    )

def create_interview_question_generator_agent():
    """Create an improved interview question generator agent."""
    llm = LLM(
    model="gemini/gemini-2.0-flash-lite",  
    temperature=0.2,
    api_key=st.secrets["GEMINI_API_KEY"] 
    )
    return Agent(
        role="Technical Interview Question Generator",
        goal="Create tailored technical interview questions based on candidate's profile and identified gaps",
        backstory="""You are an expert at crafting technical interview questions that explore the depths of a candidate's knowledge.
        Your questions help reveal actual understanding rather than memorized answers by targeting specific technologies and
        projects the candidate has worked with. You know how to design questions that assess both technical competence and
        problem-solving approach. You create a mix of questions that verify claimed expertise, explore potential skill gaps,
        and evaluate the candidate's ability to apply their knowledge in real-world scenarios. Your questions go beyond basics
        to probe depth of understanding while remaining fair and relevant to the role.""",
        verbose=True,
        allow_delegation=True,
        tools=[SearchTool()],
        llm=llm
    )

def create_evaluation_pipeline_agent():
    """Create a new agent for implementing the stage-by-stage evaluation pipeline."""
    llm = LLM(
    model="gemini/gemini-2.0-flash-lite",  
    temperature=0.2,
    api_key=st.secrets["GEMINI_API_KEY"] 
    )
    return Agent(
        role="Evaluation Pipeline Manager",
        goal="Implement a comprehensive stage-by-stage evaluation process and produce a final numerical score",
        backstory="""You are a specialist in systematic candidate evaluation processes. You implement a structured
        evaluation pipeline that breaks down candidate assessment into distinct stages, each with clear criteria
        and scoring mechanisms. You understand how to weigh different aspects of a candidate's profile according
        to their relevance to the job. You produce objective, reproducible scoring that helps reduce bias in the
        hiring process while identifying candidates with the highest potential for success in the role.""",
        verbose=True,
        allow_delegation=True,
        tools=[],
        llm=llm
    )
def create_summary_generator_agent():
    """Create an improved summary generator agent."""
    llm = LLM(
    model="gemini/gemini-2.0-flash-lite",  
    temperature=0.2,
    api_key=st.secrets["GEMINI_API_KEY"] 
    )
    return Agent(
        role="Assessment Summary Specialist",
        goal="Create engaging, concise, and visually appealing assessment reports for hiring managers",
        backstory="""You excel at synthesizing complex technical evaluations into clear, actionable, and visually engaging summaries.
        Your reports highlight candidate strengths, potential areas for growth, and relevance to specific roles
        in ways that help hiring managers make informed decisions. You have a talent for presenting information
        clearly and concisely, using visual elements like icons and rating scales to make reports easy to scan
        while still being comprehensive. You know how to create reports that are both professional and engaging,
        focusing on the most relevant insights while avoiding unnecessary details.""",
        verbose=True,
        allow_delegation=True,
        tools=[],
        llm=llm
    )

def create_parse_resume_task(document_parser_agent, resume_text):
    """Create an improved resume parsing task with focus on profile link extraction."""
    return Task(
        description=f"""
        Analyze the following resume text and extract structured information including:
        1. Personal information (name, contact details)
        2. Skills (programming languages, frameworks, tools)
        3. Work experience (companies, roles, responsibilities, achievements)
        4. Education (degrees, institutions, graduation dates)
        5. Projects (descriptions, technologies used)
        6. ALL public coding profile links (GitHub, LeetCode, Stack Overflow, etc.)
        7. ALL certification links, certification names, and credential IDs
        
        IMPORTANT: Search carefully for URLs and usernames that might indicate profiles on:
        - GitHub (look for github.com URLs or GitHub username mentions)
        - LeetCode (look for leetcode.com URLs or LeetCode username mentions)
        - Any certification platforms (Coursera, Udemy, AWS, Microsoft, etc.)
        
        Even if links are not explicitly labeled as profiles, extract any URLs or username mentions
        that might lead to coding profiles or certifications. Consider different text formats and
        look anywhere in the document, including headers, footers, and contact sections.
        
        Resume Text:
        {resume_text}
        
        Return the information in a structured JSON format with a dedicated "profiles" section
        and "certifications" section.
        """,
        agent=document_parser_agent,
        expected_output="A structured JSON with parsed resume information, identified coding profiles, and certifications"
    )

def create_scrape_profiles_task(profile_scraper_agent, profile_links):
    """Create an improved profile scraping task."""
    return Task(
        description=f"""
        Collect detailed information from the following coding profiles:
        {json.dumps(profile_links, indent=2)}
        
        For each profile:
        1. Gather comprehensive data about the candidate's activity
        2. Identify key technical metrics (languages used, projects, contributions, problem-solving ability)
        3. Look for evidence of skills claimed in their resume
        4. Analyze code quality, project complexity, and development patterns
        5. Determine the candidate's activity level and consistency of contributions
        6. Identify collaborative behaviors (such as pull requests, code reviews, etc.)
        
        For GitHub:
        - Analyze repository quality beyond just counting them
        - Examine commit patterns and contribution history
        - Check for meaningful projects vs forks or tutorial code
        - Look at code complexity and quality when possible
        
        For LeetCode:
        - Assess problem difficulty distribution (easy/medium/hard)
        - Identify algorithmic strengths and weaknesses
        - Analyze solution quality if available
        
        Return the collected information in a structured JSON format with separate sections for each platform.
        """,
        agent=profile_scraper_agent,
        expected_output="A structured JSON with detailed profile data from each platform"
    )


def create_evaluate_skills_task(technical_evaluator_agent, resume_data, profile_data, certification_data=None):
    """Create an improved skill evaluation task."""
    # Extract certifications directly from resume_data if certification_data is None
    certifications_to_consider = certification_data if certification_data else resume_data.get("certifications", {})
    
    return Task(
        description=f"""
        Evaluate the candidate's technical skills based on:
        
        Resume Information:
        {json.dumps(resume_data, indent=2)}
        
        Profile Data:
        {json.dumps(profile_data, indent=2)}
        
        Consider these certifications (without validation):
        {json.dumps(certifications_to_consider, indent=2)}
        
        Provide an objective assessment of:
        1. Technical skill levels in each language/framework/tool with a 1-10 rating
        2. Evidence of practical application of claimed skills in projects, certifications or experience
        3. Depth of knowledge in primary areas
        4. Learning ability and adaptability based on skill acquisition timeline
        5. Areas of technical strength with specific evidence from resume/profiles/certifications
        6. Potential growth areas or skill gaps
        7. Specialization vs. generalist assessment
        8. Evaluate coding skills based coding platforms data such as leetcode profile data
        9. Evalaute open source contribuion levels based on girhub profile data
        10. Evaluate active participation of the candidate based on linkedIn profile data
        11. Evaluate achievements of the candidate 
        12. Evaluate industry exposure, academic achievemnts or research paublications of the candidate
        13. Evaluate Novelty of the projects done by candidate either in interships or jobs or academic projects
        
        For each skill:
        - Distinguish between claimed skills and demonstrated skills
        - Consider both experience (years) and evidence of mastery
        - Weigh profile evidence more heavily than resume claims
        - Take into account certifications at face value (assume they are valid)
        
        Return your evaluation in a structured JSON format with skills grouped by category
        (languages, frameworks, tools, etc.) and each assigned a numerical rating.
        """,
        agent=technical_evaluator_agent,
        expected_output="A structured JSON with detailed technical skill evaluation"
    )

def create_job_match_task(job_alignment_agent, skill_evaluation, job_requirements):
    """Create an improved job match task."""
    return Task(
        description=f"""
        Determine how well the candidate matches the following job requirements:
        {job_requirements}
        
        Using their skill evaluation:
        {json.dumps(skill_evaluation, indent=2)}
        
        Assess:
        1. Essential requirements match (must-have skills)
        2. Preferred requirements match (nice-to-have skills)
        3. Experience level alignment
        4. Cultural fit indicators
        5. Growth potential for the role
        
        For each requirement:
        - Identify if it's satisfied directly or through a transferable skill
        - Consider not just the presence of a skill but its rated level
        - Note where the candidate exceeds requirements
        - Note where the candidate falls short but may be able to learn
        - Note where critical gaps exist
        
        Provide a match percentage (0-100%) and detailed explanation for each major requirement.
        Also calculate an overall match percentage that weights essential requirements more heavily.
        Return your assessment in a structured JSON format.
        """,
        agent=job_alignment_agent,
        expected_output="A structured JSON with detailed job requirement matches"
    )

def create_evaluation_pipeline_task(evaluation_pipeline_agent, resume_data, profile_data, certification_data, skill_evaluation, job_match):
    """Create a new task for the stage-by-stage evaluation pipeline."""
    return Task(
        description=f"""
        Implement a comprehensive evaluation pipeline with the following stages, using:
        
        Resume Information:
        {json.dumps(resume_data, indent=2)}
        
        Profile Data:
        {json.dumps(profile_data, indent=2)}
        
        Certification Data:
        {json.dumps(certification_data, indent=2)}
        
        Skill Evaluation:
        {json.dumps(skill_evaluation, indent=2)}
        
        Job Match Assessment:
        {json.dumps(job_match, indent=2)}
        
        **Stage 1: Basic Eligibility Check** (20 points)
        - Extract and verify: education, years of experience, certifications, location (if required)
        - Calculate points based on minimum requirements being met
        
        **Stage 2: Skill Match Evaluation** (40 points)
        - Use the skill evaluation and job match data
        - Classify: Matched Skills, Missing Skills, Bonus Skills
        - Weight essential skills higher than preferred skills
        
        **Stage 3: Domain-Specific Experience Check** (15 points)
        - Identify domain-related keywords from the job requirements
        - Score based on evidence of domain experience in resume and profiles
        
        **Stage 4: Role Fit and Achievements Alignment** (25 points)
        - Assess alignment of resume achievements with job responsibilities
        - Prioritize impact-driven achievements and leadership experience if relevant
        
        For each stage:
        - Provide a numeric score
        - Include a justification for the score
        - Highlight key factors that influenced the score
        
        Finally:
        - Calculate a total score out of 100
        - Provide a final recommendation: Proceed to interview / Hold / Reject
          (Proceed: 70+, Hold: 50-69, Reject: <50)
        
        Return a structured JSON with results for each stage, total score, and recommendation.
        """,
        agent=evaluation_pipeline_agent,
        expected_output="A structured JSON containing scores and evaluations for each stage of the pipeline"
    )


def create_generate_questions_task(interview_question_generator, resume_data, profile_data, skill_evaluation, job_match):
    """Create an improved interview question generation task."""
    return Task(
        description=f"""
        Generate strategic technical interview questions based on:
        
        Resume Information:
        {json.dumps(resume_data, indent=2)}
        
        Profile Data:
        {json.dumps(profile_data, indent=2)}
        
        Skill Evaluation:
        {json.dumps(skill_evaluation, indent=2)}
        
        Job Match Assessment:
        {json.dumps(job_match, indent=2)}
        
        Create 5-7 tailored questions that:
        1. Verify depth of knowledge in claimed expertise areas
        2. Explore specific projects mentioned in resume or profiles
        3. Address identified skill gaps relevant to the job requirements
        4. Include at least one system design question related to their experience
        5. Include at least one problem-solving question that relates to their domain
        6. Test adaptability by asking about unfamiliar but related technologies
        
        For each question:
        - Make it specific to this candidate (reference specific projects or experience)
        - Clearly state what you're assessing with this question
        - Provide guidance on what a good answer would include
        - Structure questions that reveal thinking process, not just knowledge
        - Include difficulty level (Basic, Intermediate, Advanced)
        
        Return the questions in a structured JSON format.
        """,
        agent=interview_question_generator,
        expected_output="A structured JSON with tailored technical interview questions"
    )

def create_summary_task(summary_generator, resume_data, profile_data, skill_evaluation, job_match, interview_questions, evaluation_results):
    """Create an improved summary generation task with focus on readability and visual appeal."""
    return Task(
        description=f"""
        Create a comprehensive,compelling, visually engaging, yet concise assessment summary based on:
        
        Resume Information:
        {json.dumps(resume_data, indent=2)}
        
        Profile Data:
        {json.dumps(profile_data, indent=2)}
        
        Skill Evaluation:
        {json.dumps(skill_evaluation, indent=2)}
        
        Job Match Assessment:
        {json.dumps(job_match, indent=2)}
        
        Suggested Interview Questions:
        {json.dumps(interview_questions, indent=2)}
        
        Evaluation Results:
        {json.dumps(evaluation_results, indent=2)}
        
        Create a summary that:
        1. A candidate overview highlighting their experience and education
        2. Includes a visually appealing score card with the overall evaluation score
        3. Key strengths with evidence from their resume and profiles
        4. Potential areas for growth or skill gaps
        5. Summarizes overall job match with a clear recommendation
        6. Recommendations for the hiring team
        7. Suggested next steps
        8. Uses visually engaging elements like:
           - Star ratings (★★★★☆) for key skills
           - Emojis as bullet points for visual scanning
           - Rating scales for key metrics
           - Clear headings and subheadings
        
        The summary should be structured for easy reading by hiring managers. It should be readable within 1 min and provide valuble insights.
        Format in visually appealing markdown, optimized for a hiring manager who has limited time.
        Return a well-formatted markdown document
        """,
        agent=summary_generator,
        expected_output="A concise, visually engaging markdown assessment summary"
    )


def run_skill_assessment(resume_file, job_requirements):
    print("Starting enhanced skill assessment process...")
    
    # Step 1: Extract text from resume document
    print("Extracting text from resume...")
    resume_text = extract_text_from_document(resume_file)
    
    # Step 2: Create agents
    print("Creating specialized agents...")
    document_parser_agent = create_document_parser_agent()
    profile_scraper_agent = create_profile_scraper_agent()  
    technical_evaluator_agent = create_technical_evaluator_agent()
    job_alignment_agent = create_job_alignment_agent()
    evaluation_pipeline_agent = create_evaluation_pipeline_agent()  
    interview_question_generator = create_interview_question_generator_agent()
    summary_generator_agent = create_summary_generator_agent()
    
    # Step 3: Define tasks and create parsing crew
    print("Creating assessment tasks...")
    parse_resume_task = create_parse_resume_task(document_parser_agent, resume_text)
    
    # Step 4: Create crew and execute
    print("Executing document parsing...")
    parsing_crew = Crew(
        agents=[document_parser_agent],
        tasks=[parse_resume_task],
        verbose=True,
        process=Process.sequential
    )
    
    resume_result = parsing_crew.kickoff()
    print("\nResume parsing complete!")
    
    try:
        # Access the actual string value from the CrewOutput object
        resume_result_str = str(resume_result)
        resume_data = json.loads(resume_result_str)
    except json.JSONDecodeError:
        # Try to extract JSON from text if the result isn't directly parseable
        try:
            json_match = re.search(r'```json\n(.*?)\n```', resume_result_str, re.DOTALL)
            if json_match:
                resume_data = json.loads(json_match.group(1))
            else:
                print("Warning: Could not parse resume result as JSON. Using raw text.")
                resume_data = {"raw_result": resume_result_str}
        except Exception as e:
            print(f"Error parsing resume JSON: {str(e)}")
            resume_data = {"raw_result": resume_result_str}
    
    # Step 5: Extract profile links and scrape profiles
    print("\nExtracting profile links...")
    if isinstance(resume_data, dict) and "raw_result" not in resume_data:
        # Use the structured data if available
        profile_links = resume_data.get("profiles", {})
        if not profile_links:
            # Try to find links in the structured data
            profile_links = extract_profile_links(str(resume_data))
    else:
        # Use the raw text for extraction
        profile_links = extract_profile_links(resume_text)
    
    print(f"Found profile links: {profile_links}")
    
    if profile_links:
        print("\nScraping coding profiles...")
        scrape_profiles_task = create_scrape_profiles_task(profile_scraper_agent, profile_links)
        
        profile_crew = Crew(
            agents=[profile_scraper_agent],
            tasks=[scrape_profiles_task],
            verbose=True,
            process=Process.sequential
        )
        
        profile_result = profile_crew.kickoff()
        print("\nProfile scraping complete!")
        
        try:
            # Convert CrewOutput to string before parsing as JSON
            profile_result_str = str(profile_result)
            profile_data = json.loads(profile_result_str)
        except json.JSONDecodeError:
            # Try to extract JSON from text
            try:
                json_match = re.search(r'```json\n(.*?)\n```', profile_result_str, re.DOTALL)
                if json_match:
                    profile_data = json.loads(json_match.group(1))
                else:
                    print("Warning: Could not parse profile result as JSON. Using raw text.")
                    profile_data = {"raw_result": profile_result_str}
            except Exception as e:
                print(f"Error parsing profile JSON: {str(e)}")
                profile_data = {"raw_result": profile_result_str}
    else:
        print("No coding profiles found. Proceeding with resume data only.")
        profile_data = {}
    
    # Step 6: Verify certifications if present
    print("\nExtracting certification information...")
    if isinstance(resume_data, dict) and "raw_result" not in resume_data:
        # Use the structured data if available
        certifications = resume_data.get("certifications", {})
        if not certifications:
            # Try to find certifications in the structured data
            certifications = extract_certifications(str(resume_data))
    else:
        # Use the raw text for extraction
        certifications = extract_certifications(resume_text)
    
    print(f"Found certifications: {certifications}")
    
    # Step 7: Evaluate technical skills
    print("\nEvaluating technical skills...")
    evaluate_skills_task = create_evaluate_skills_task(technical_evaluator_agent, resume_data, profile_data, certifications)
    
    skills_crew = Crew(
        agents=[technical_evaluator_agent],
        tasks=[evaluate_skills_task],
        verbose=True,
        process=Process.sequential
    )
    
    skills_result = skills_crew.kickoff()
    print("\nSkill evaluation complete!")
    
    try:
        # Convert CrewOutput to string before parsing as JSON
        skills_result_str = str(skills_result)
        skill_evaluation = json.loads(skills_result_str)
    except json.JSONDecodeError:
        # Try to extract JSON from text
        try:
            json_match = re.search(r'```json\n(.*?)\n```', skills_result_str, re.DOTALL)
            if json_match:
                skill_evaluation = json.loads(json_match.group(1))
            else:
                print("Warning: Could not parse skill evaluation as JSON. Using raw text.")
                skill_evaluation = {"raw_result": skills_result_str}
        except Exception as e:
            print(f"Error parsing skill evaluation JSON: {str(e)}")
            skill_evaluation = {"raw_result": skills_result_str}
    
    # Step 8: Assess job match
    print("\nAssessing job match...")
    job_match_task = create_job_match_task(job_alignment_agent, skill_evaluation, job_requirements)
    
    job_match_crew = Crew(
        agents=[job_alignment_agent],
        tasks=[job_match_task],
        verbose=True,
        process=Process.sequential
    )
    
    job_match_result = job_match_crew.kickoff()
    print("\nJob match assessment complete!")
    
    try:
        # Convert CrewOutput to string before parsing as JSON
        job_match_result_str = str(job_match_result)
        job_match = json.loads(job_match_result_str)
    except json.JSONDecodeError:
        # Try to extract JSON from text
        try:
            json_match = re.search(r'```json\n(.*?)\n```', job_match_result_str, re.DOTALL)
            if json_match:
                job_match = json.loads(json_match.group(1))
            else:
                print("Warning: Could not parse job match as JSON. Using raw text.")
                job_match = {"raw_result": job_match_result_str}
        except Exception as e:
            print(f"Error parsing job match JSON: {str(e)}")
            job_match = {"raw_result": job_match_result_str}
    
    # Step 9: Run the evaluation pipeline
    print("\nRunning evaluation pipeline...")
    evaluation_pipeline_task = create_evaluation_pipeline_task(
        evaluation_pipeline_agent,
        resume_data,
        profile_data,
        certifications,
        skill_evaluation,
        job_match
    )
    
    evaluation_crew = Crew(
        agents=[evaluation_pipeline_agent],
        tasks=[evaluation_pipeline_task],
        verbose=True,
        process=Process.sequential
    )
    
    evaluation_result = evaluation_crew.kickoff()
    print("\nEvaluation pipeline complete!")
    
    try:
        evaluation_result_str = str(evaluation_result)
        evaluation_results = json.loads(evaluation_result_str)
    except json.JSONDecodeError:
        try:
            json_match = re.search(r'```json\n(.*?)\n```', evaluation_result_str, re.DOTALL)
            if json_match:
                evaluation_results = json.loads(json_match.group(1))
            else:
                print("Warning: Could not parse evaluation results as JSON. Using raw text.")
                evaluation_results = {"raw_result": evaluation_result_str}
        except Exception as e:
            print(f"Error parsing evaluation results JSON: {str(e)}")
            evaluation_results = {"raw_result": evaluation_result_str}
    
    # Step 10: Generate interview questions
    print("\nGenerating interview questions...")
    questions_task = create_generate_questions_task(
        interview_question_generator, 
        resume_data, 
        profile_data,
        skill_evaluation,
        job_match
    )
    
    questions_crew = Crew(
        agents=[interview_question_generator],
        tasks=[questions_task],
        verbose=True,
        process=Process.sequential
    )
    
    questions_result = questions_crew.kickoff()
    print("\nInterview questions generated!")
    
    try:
        # Convert CrewOutput to string before parsing as JSON
        questions_result_str = str(questions_result)
        interview_questions = json.loads(questions_result_str)
    except json.JSONDecodeError:
        # Try to extract JSON from text
        try:
            json_match = re.search(r'```json\n(.*?)\n```', questions_result_str, re.DOTALL)
            if json_match:
                interview_questions = json.loads(json_match.group(1))
            else:
                print("Warning: Could not parse interview questions as JSON. Using raw text.")
                interview_questions = {"raw_result": questions_result_str}
        except Exception as e:
            print(f"Error parsing interview questions JSON: {str(e)}")
            interview_questions = {"raw_result": questions_result_str}
    
    # Step 11: Generate final summary
    print("\nGenerating final assessment summary...")
    summary_task = create_summary_task(
        summary_generator_agent,
        resume_data,
        profile_data,
        skill_evaluation,
        job_match,
        interview_questions,
        evaluation_results
    )
    
    summary_crew = Crew(
        agents=[summary_generator_agent],
        tasks=[summary_task],
        verbose=True,
        process=Process.sequential
    )
    
    summary_result = summary_crew.kickoff()
    print("\nFinal assessment summary generated!")
    
    # Step 12: Return all results
    results = {
        "resume_data": resume_data,
        "profile_data": profile_data,
        "certification_data": certifications,
        "skill_evaluation": skill_evaluation,
        "job_match": job_match,
        "evaluation_results": evaluation_results,
        "interview_questions": interview_questions,
        "summary": summary_result
    }
    
    return results


# Create sample job requirements
# Replace the last part of skill_ass.py with this:

sample_job_requirements = """
Job Title: Senior Full Stack Developer

Essential Requirements:
- 5+ years of experience in software development
- Strong proficiency in Python and JavaScript
- Experience with React.js and Django
- Knowledge of RESTful API design and implementation
- Experience with cloud platforms (AWS or Azure)
- Understanding of CI/CD pipelines and DevOps practices

Preferred Skills:
- Experience with TypeScript and Node.js
- Knowledge of container technologies (Docker, Kubernetes)
- Experience with database design and optimization
- Familiarity with microservices architecture
- Experience mentoring junior developers

Responsibilities:
- Design, develop, and maintain web applications
- Collaborate with cross-functional teams to define and implement new features
- Write clean, maintainable, and efficient code
- Participate in code reviews and provide constructive feedback
- Troubleshoot production issues and optimize application performance
- Stay up-to-date with emerging trends and technologies
"""

# Keep the make_json_serializable function as it's used elsewhere
def make_json_serializable(obj):
    """Convert any non-serializable objects to strings"""
    if hasattr(obj, '__class__') and obj.__class__.__name__ == 'CrewOutput':
        return str(obj)
    elif isinstance(obj, dict):
        return {k: make_json_serializable(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [make_json_serializable(item) for item in obj]
    else:
        return obj
    


# with open("Nandini_kuppala_latest.pdf", "rb") as file:
#     results = run_skill_assessment(file, sample_job_requirements)
# 
# serializable_results = make_json_serializable(results)
# with open("assessment_results.json", "w") as f:
#     json.dump(serializable_results, f, indent=2)
